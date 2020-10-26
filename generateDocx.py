# -*- coding: utf-8 -*-
import os
import requests
from docx import Document
from docx.shared import Cm
from pdf2png import convertPdfIntoPng
from pngSplit import cutPngIntoGrid


def getTxtContent(readFile, writeFile, geoType):
    r = requests.get(readFile).content
    contents = [c.decode('utf-8') for c in r.split("\r\n")]

    d = Document("template_report.docx")

    d.add_heading(text=u"商户基本信息", level=1)
    d.add_heading(text=u"商户区域分布", level=2)
    d.add_paragraph(text=contents[0])
    d.add_paragraph(text=u"""表 1-1 烟草商户区域分布""", style="Caption")
    if geoType == "city":
        d.add_picture("test_0.png", width=Cm(13.0))
    else:
        d.add_picture("test_10.png", width=Cm(13.0))

    d.add_heading(text=u"商户经营者属性统计", level=2)
    d.add_heading(text=u"经营者性别分布", level=3)
    d.add_paragraph(text=contents[1])
    d.add_paragraph(text=u"""图 1-1 烟草商户性别分布""", style="Caption")
    d.add_picture("test_1.png", width=Cm(13.0))

    d.add_heading(text=u"经营者年龄分布", level=3)
    d.add_paragraph(text=contents[2])
    d.add_paragraph(text=u"""图 1-2 烟草商户年龄分布""", style="Caption")
    d.add_picture("test_2.png", width=Cm(13.0))

    d.add_heading(text=u"商户店铺属性统计", level=2)
    d.add_heading(text=u"店铺城乡类型分布", level=3)
    d.add_paragraph(text=contents[3])
    if contents[3].find(u"没有乡村") == -1:
        d.add_paragraph(text=u"""图 1-3 烟草商户城乡分布""", style="Caption")
        d.add_picture("test_3.png", width=Cm(13.0))

    d.add_heading(text=u"店铺业态分布", level=3)
    d.add_paragraph(text=contents[4])
    d.add_paragraph(text=u"""图 1-4 烟草商户业态分布""", style="Caption")
    d.add_picture("test_4.png", width=Cm(13.0))

    d.add_heading(text=u"店铺权属分布", level=3)
    d.add_paragraph(text=contents[5])
    d.add_paragraph(text=u"""图 1-5 烟草商户店铺权属分布""", style="Caption")
    d.add_picture("test_5.png", width=Cm(13.0))

    d.add_heading(text=u"店铺面积分布", level=3)
    d.add_paragraph(text=contents[6])
    d.add_paragraph(text=u"""图 1-6 烟草商户店铺面积分布""", style="Caption")
    d.add_picture("test_6.png", width=Cm(13.0))

    d.add_heading(text=u"商户经营信息", level=1)
    d.add_heading(text=u"商户经营时长分布", level=2)
    d.add_paragraph(text=contents[7])
    d.add_paragraph(text=u"""图 2-1 烟草商户经营时长分布""", style="Caption")
    d.add_picture("test_7.png", width=Cm(13.0))

    d.add_heading(text=u"商户平均订烟周期分布", level=2)
    d.add_paragraph(text=contents[8])
    d.add_paragraph(text=u"""图 2-2 烟草商户平均订烟周期分布""", style="Caption")
    d.add_picture("test_8.png", width=Cm(13.0))

    d.add_heading(text=u"商户月均订烟金额分布", level=2)
    d.add_paragraph(text=contents[9])
    d.add_paragraph(text=u"""图 2-3 烟草商户月均订烟金额分布""", style="Caption")
    d.add_picture("test_9.png", width=Cm(13.0))

    d.save(writeFile)


def generateDocx(reportPlace, geoType="city"):
    convertPdfIntoPng("exportPdf.pdf", "test.png")
    cutPngIntoGrid("test.png", 12, 2, [x * 2 + 1 for x in range(11)])
    getTxtContent("cityPortrait.txt", reportPlace + u"商户画像报告.docx", geoType=geoType)
    os.remove("E:\\Download\\exportPdf.pdf")
    print(reportPlace + u"商户画像报告.docx已生成。")


if __name__ == "__main__":
    generateDocx(u"辽宁省", geoType="province")